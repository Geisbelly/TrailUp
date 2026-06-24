from io import BytesIO

from docx import Document


def gerar_docx(
    *,
    titulo: str,
    resumo: str,
    secoes: list[str],
) -> bytes:
    doc = Document()
    doc.add_heading(titulo or "Documento personalizado", level=1)
    if resumo:
        doc.add_paragraph(resumo)
    for index, secao in enumerate(secoes, start=1):
        doc.add_heading(f"Seção {index}", level=2)
        doc.add_paragraph(str(secao))
    output = BytesIO()
    doc.save(output)
    return output.getvalue()
