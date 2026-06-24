"""Extrator de DOCX: python-docx → lista de Block."""

from __future__ import annotations

import io
import re

from app.ingestion.models import Block, BlockKind, FileFamily


def extract(raw: bytes, *, filename: str = "documento.docx") -> dict:
    """
    Extrai blocos estruturados de um DOCX.

    Retorna dict com:
      - family, title, blocks (list[Block]), metadata
    """
    try:
        from docx import Document
    except ImportError:
        return _empty(filename)

    try:
        doc = Document(io.BytesIO(raw))
    except Exception:
        return _empty(filename)

    blocks: list[Block] = []
    block_counter = 0
    title = filename.replace(".docx", "").replace("_", " ").replace("-", " ").title()
    first_heading_found = False

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = (para.style.name or "").lower()
        block_counter += 1

        if "heading 1" in style_name or style_name == "title":
            if not first_heading_found:
                title = text
                first_heading_found = True
            blocks.append(
                Block(
                    block_id=f"p{block_counter}",
                    kind=BlockKind.HEADING,
                    text=text,
                    order=block_counter,
                    source_ref={"paragraph": block_counter},
                    metadata={"level": 1, "style": para.style.name},
                )
            )
        elif "heading 2" in style_name:
            blocks.append(
                Block(
                    block_id=f"p{block_counter}",
                    kind=BlockKind.SUBHEADING,
                    text=text,
                    order=block_counter,
                    source_ref={"paragraph": block_counter},
                    metadata={"level": 2, "style": para.style.name},
                )
            )
        elif "heading" in style_name:
            blocks.append(
                Block(
                    block_id=f"p{block_counter}",
                    kind=BlockKind.SUBHEADING,
                    text=text,
                    order=block_counter,
                    source_ref={"paragraph": block_counter},
                    metadata={"level": 3, "style": para.style.name},
                )
            )
        elif "list" in style_name or re.match(r"^[\-\*\•\·]\s", text) or re.match(r"^\d+[\.\)]\s", text):
            blocks.append(
                Block(
                    block_id=f"p{block_counter}",
                    kind=BlockKind.LIST_ITEM,
                    text=text,
                    order=block_counter,
                    source_ref={"paragraph": block_counter},
                    metadata={},
                )
            )
        elif "quote" in style_name or "block" in style_name:
            blocks.append(
                Block(
                    block_id=f"p{block_counter}",
                    kind=BlockKind.QUOTE,
                    text=text,
                    order=block_counter,
                    source_ref={"paragraph": block_counter},
                    metadata={},
                )
            )
        else:
            blocks.append(
                Block(
                    block_id=f"p{block_counter}",
                    kind=BlockKind.PARAGRAPH,
                    text=text,
                    order=block_counter,
                    source_ref={"paragraph": block_counter},
                    metadata={},
                )
            )

    # Extrair tabelas como texto
    for tbl_idx, table in enumerate(doc.tables, start=1):
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            block_counter += 1
            blocks.append(
                Block(
                    block_id=f"table_{tbl_idx}",
                    kind=BlockKind.TABLE,
                    text="\n".join(rows),
                    order=block_counter,
                    source_ref={"table": tbl_idx},
                    metadata={"rows": len(rows)},
                )
            )

    return {
        "family": FileFamily.TEXT,
        "title": title,
        "blocks": blocks,
        "metadata": {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "filename": filename,
        },
    }


def _empty(filename: str) -> dict:
    return {
        "family": FileFamily.TEXT,
        "title": filename,
        "blocks": [],
        "metadata": {"filename": filename, "error": "extraction_failed"},
    }
